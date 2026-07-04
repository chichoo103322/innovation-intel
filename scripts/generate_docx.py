#!/usr/bin/env python3
"""
《创新常州·对标快讯》Word 文档生成器
严格对齐参考格式：►标题（日期）：摘要。创新洞察：分析。
用法: python3 generate_docx.py --type daily
"""
import sys
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

PROJECT_DIR = Path(__file__).parent.parent
ISSUE_FILE = PROJECT_DIR / "cache" / "issue_counter.json"


def get_issue_numbers() -> tuple:
    """
    获取当前期号。
    如果 settings.yaml 中配置了 first_issue_date，按工作日自动计算期号；
    否则使用计数器自动递增。
    """
    # 尝试读取配置文件中的起始日期
    from run_daily import load_config
    try:
        cfg = load_config()
        first_date_str = cfg.get("first_issue_date", "")
    except Exception:
        first_date_str = ""

    if first_date_str:
        # 按工作日计算：从起始日期到今天的周一到周五天数
        from datetime import date, timedelta
        start = datetime.strptime(first_date_str, "%Y-%m-%d").date()
        today = date.today()
        if today < start:
            return 1, 1
        # 统计从 start 到 today 的工作日数
        delta = (today - start).days
        weekdays = 0
        for i in range(delta + 1):
            d = start + timedelta(days=i)
            if d.weekday() < 5:  # 周一到周五
                weekdays += 1
        return weekdays, weekdays

    # 回退：计数器模式
    ISSUE_FILE.parent.mkdir(parents=True, exist_ok=True)
    if ISSUE_FILE.exists():
        data = json.loads(ISSUE_FILE.read_text())
    else:
        data = {"issue": 0, "total": 0}
    data["issue"] += 1
    data["total"] += 1
    ISSUE_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))
    return data["issue"], data["total"]


def reset_issue_counter():
    """手动重置期号计数器"""
    ISSUE_FILE.write_text(json.dumps({"issue": 0, "total": 0}, ensure_ascii=False, indent=2))
    print("[已重置] 期号计数器归零")


# 四号 = 14pt
BODY_SIZE = Pt(14)
INDENT_2CHAR = Pt(28)  # 2个四号字符宽度


def set_run_font(run, ea_font, size=BODY_SIZE, bold=False, color=None):
    """统一设置字体：中文用 ea_font，西文/数字用 Times New Roman"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), ea_font)
    if color:
        run.font.color.rgb = color


def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)
    return doc


def add_header_block(doc, title_text, issue_no, total_no, date_cn):
    """标题区"""
    # 主标题：华文中宋 48pt 蓝色加粗
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    run = p.add_run(title_text)
    set_run_font(run, '华文中宋', Pt(48), bold=True, color=RGBColor(0x00, 0x70, 0xC0))

    # 期号行：黑体 四号，第X期数字蓝色，总第X期数字黑色 + 蓝色底线
    # issue_no 为 0 时跳过（周报/月报不需要期号）
    if issue_no > 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(2)
        parts = [
            ('2026年第', None), (str(issue_no), RGBColor(0x00, 0x70, 0xC0)),
            ('期   总第', None), (str(total_no), None),
            ('期', None),
        ]
        for text, color in parts:
            run = p.add_run(text)
            set_run_font(run, '黑体', BODY_SIZE, color=color)
        # 期号行下方蓝色底线
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '6',
            qn('w:space'): '4', qn('w:color'): '1a3c6e',
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # 日期行：黑体 四号，右对齐
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.space_after = Pt(14)
    run = p.add_run(date_cn)
    set_run_font(run, '黑体', BODY_SIZE)


def add_section_title(doc, text):
    """板块标题：黑体 三号(16pt) 加粗"""
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, '黑体', Pt(16), bold=True)


def add_news_item(doc, title, date_info, summary, insight, source="", url=""):
    """
    严格对齐参考格式（两段式）：
    ►标题（日期）：内容摘要。
    创新洞察：分析内容。（来源：XXX）—— 来源可点击跳转
    """
    # 段落1：►标题（日期）：摘要  —— 仿宋_GB2312 四号，首行缩进2字符
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    p.paragraph_format.first_line_indent = INDENT_2CHAR

    run = p.add_run('►')
    set_run_font(run, '仿宋_GB2312')

    run = p.add_run(title)
    set_run_font(run, '仿宋_GB2312', bold=True)

    if date_info:
        run = p.add_run(f'（{date_info}）')
        set_run_font(run, '仿宋_GB2312')

    run = p.add_run('：')
    set_run_font(run, '仿宋_GB2312')

    run = p.add_run(summary)
    set_run_font(run, '仿宋_GB2312')

    # 段落2：创新洞察：分析内容。 —— 楷体_GB2312 四号，首行缩进2字符
    p2 = doc.add_paragraph()
    p2.space_before = Pt(2)
    p2.space_after = Pt(2)
    p2.paragraph_format.first_line_indent = INDENT_2CHAR

    run = p2.add_run('创新洞察：')
    set_run_font(run, '楷体_GB2312', bold=True)

    run = p2.add_run(insight)
    set_run_font(run, '楷体_GB2312')

    # 段落3：信息来源链接（可点击跳转）
    # 格式对齐参考文档：｜信息来源：URL（微软雅黑 9pt 灰色，URL可点击）
    if url:
        p3 = doc.add_paragraph()
        p3.space_before = Pt(2)
        p3.space_after = Pt(6)

        # "｜信息来源：" 标签
        run = p3.add_run('｜信息来源：')
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # 可点击的超链接 URL
        r_id = doc.part.relate_to(url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True)
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        _make_font_el(rPr, Pt(9), '微软雅黑', '微软雅黑', RGBColor(0x88, 0x88, 0x88))
        r.append(rPr)

        # 超链接样式
        c = OxmlElement('w:u')
        c.set(qn('w:val'), 'single')
        c.set(qn('w:color'), '888888')
        rPr.append(c)

        t = OxmlElement('w:t')
        t.text = url
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        hl.append(r)
        p3._element.append(hl)
    elif source:
        p3 = doc.add_paragraph()
        p3.space_before = Pt(2)
        p3.space_after = Pt(6)
        run = p3.add_run(f'｜信息来源：{source}')
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _make_font_el(rPr, size, west_font, ea_font, color):
    """构建 rPr 字体元素"""
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), west_font)
    rFonts.set(qn('w:hAnsi'), west_font)
    rFonts.set(qn('w:eastAsia'), ea_font)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size.pt * 2)))
    rPr.append(sz)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), str(color))
        rPr.append(c)


def generate_daily_from_data(sections_data, output_path=None, title_text='创新常州·对标快讯'):
    """根据结构化数据生成日报（由 run_daily.py 调用）"""
    today = datetime.now()
    date_cn = today.strftime('%Y年%m月%d日')
    date_fn = today.strftime('%Y-%m-%d')
    issue, total = get_issue_numbers()

    doc = create_doc()
    add_header_block(doc, title_text, issue, total, date_cn)

    for section in sections_data:
        add_section_title(doc, f'【{section["name"]}】')
        for item in section.get("items", []):
            add_news_item(doc,
                title=item.get("title", ""),
                date_info=item.get("date", ""),
                summary=item.get("summary", ""),
                insight=item.get("insight", ""),
                source=item.get("source", ""),
                url=item.get("url", ""))

    if output_path is None:
        output_path = PROJECT_DIR / 'daily' / f'创新常州·对标快讯_{date_fn}.docx'
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'[完成] 日报已生成: {output_path}')
    return output_path


def generate_daily(output_path=None):
    """生成日报（使用示例数据，向后兼容）"""
    sample_data = [
        {
            "name": "各地科技委动态",
            "items": [
                {
                    "title": "江苏省委科技委员会全体会议召开",
                    "date": "2026.7.1",
                    "summary": "省委书记信长星主持，省长刘小涛出席。会议强调高效协同推进上海（长三角）国际科技创新中心建设，南京、苏州扛起关键支点重任，无锡、常州、南通建设高能级创新型城市。会议讨论了《江苏省人工智能驱动科学研究实施方案》，提出推动科技创新与产业创新深度融合，打造发展新质生产力重要阵地。",
                    "insight": "常州被明确列为「高能级创新型城市」建设对象，应趁势主动对接省AI驱动科研方案，争取省级示范项目落地，同时加快制定常州版高能级创新型城市建设行动方案，明确在沿沪宁产业创新带中的差异化功能定位。",
                    "source": "江苏省人民政府"
                },
                {
                    "title": "江西省委科技委第七次全体会议召开",
                    "date": "2026.7.2",
                    "summary": "省长叶建春主持，部署加快推进「十五五」科技创新规划编制，实施「2030启航计划」和「2030先锋工程」，推动人工智能重构科研范式，构建「政府引导、市场主导、社会参与」的多元化科技投入格局，聚焦突出矛盾开展靶向攻坚。",
                    "insight": "江西的「靶向攻坚」方法论和多元化科技投入格局值得借鉴。常州可在「三名联动」框架下探索「财政引导+企业主导+社会资本跟投」三级投入模式，集中资源突破新能源、合成生物等优势领域的关键瓶颈。",
                    "source": "江西新闻网"
                },
                {
                    "title": "广东省委科技委发布脑机接口产业协同发展行动计划",
                    "date": "2026.7.2",
                    "summary": "印发《广东省脑机接口科技与产业协同发展行动计划（2026-2030年）》，目标到2030年核心产业规模达百亿级，辐射上下游达千亿级。这是全国首个由省级科技委直接发布的未来产业行动计划，标志科技委正从协调机构向产业战略策源机构转型。",
                    "insight": "广东科技委直接发布产业行动计划的做法表明科技委角色正在升级。常州可参照此模式，由市委科技委直接发布新能源、合成生物等优势产业的技术创新路线图，提升科技委的战略引领力和产业话语权。",
                    "source": "广东省科技厅"
                },
            ]
        },
        {
            "name": "上海（长三角）国创中心资讯",
            "items": [
                {
                    "title": "长三角市场监管一体化发布助力国创中心建设19条举措",
                    "date": "2026.7.3",
                    "summary": "三省一市联合发布《若干举措》，涵盖科技人才「互聘共享」和资质跨区域互认、元宇宙/AI/6G等前沿领域快速获权服务、融入「一中心五支点、一廊两带」主干网等七大板块。目前长三角经营主体达3828万户，R&D经费占全国约30.55%，万人高价值发明专利29.36件。",
                    "insight": "19条举措中，「人才互聘共享」和「资质跨区域互认」对常州是直接利好，可借此机制柔性引进沪宁领军人才，降低人才流动制度壁垒。建议市科技局牵头梳理企业人才需求清单，主动对接长三角人才共享平台。",
                    "source": "澎湃新闻"
                },
                {
                    "title": "学术论坛热议国际科创中心「扩围提质」：刘庆提「没有围墙的创新中心」",
                    "date": "2026.7.2",
                    "summary": "第23届上海市社科界学术年会论坛在同济大学举行。长三角国创中心主任刘庆提出打造「没有围墙的创新中心」，推动从「问题清单」迈向「跨域规则共识」。上海全球城市研究院院长阮青提出「前研后转」分工模式——上海做前沿研发，周边城市做产业转化。",
                    "insight": "「前研后转」分工模式对常州定位意义重大——上海做前沿研发，常州做产业转化和中试放大。应围绕此定位建强中以常州创新园、科教城等转化载体，梳理与上海高校院所的合作项目，形成「沪研常转」标杆案例。",
                    "source": "网易新闻"
                },
            ]
        },
        {
            "name": "科创政策速览",
            "items": [
                {
                    "title": "武汉光谷发布智能体经济「路线图」：三年投10亿、年内万P算力",
                    "date": "2026.7.3",
                    "summary": "光谷未来三年投入超10亿元，培育100家智能体创新企业、落地1000个创新产品、汇聚10000名开发者。已建成超5000P智能算力，年内提升至10000P，企业租用算力享最高50%补贴。设立智能体场景发布厅，标杆场景最高300万元支持，百亿级AI基金集群重点投向智能体赛道。",
                    "insight": "光谷在智能体经济上的投入规模值得关注。常州新能源产业场景丰富，有条件打造「AI+新能源」特色智能体场景——如智能电网调度智能体、电池健康管理智能体等，与光谷形成差异化竞争。",
                    "source": "财联社"
                },
                {
                    "title": "南京推进AI立法：审议AI产业发展条例，大模型备案奖励20万元",
                    "date": "2026.7.3",
                    "summary": "南京首次审议《南京市人工智能产业发展和创新应用促进条例（草案）》，拟设AI「一人公司」扶持机制，出台「人工智能+软件」发展行动方案，推行「算力券」补贴机制，大模型通过备案奖励20万元，AI重点项目最高补助200万元。AI关联产业规模已突破3000亿元。",
                    "insight": "南京「算力券」机制值得常州参考。常州中小企业多，大模型训练成本是痛点，可探索设立市级「算力券」补贴，降低中小企业AI应用门槛。同时密切关注南京AI条例的立法经验，为常州未来AI产业规范化发展储备政策工具。",
                    "source": "江苏经济报"
                },
                {
                    "title": "郑州入局万亿具身智能赛道，与宇树科技共建实训创新中心",
                    "date": "2026.7.3",
                    "summary": "OpenLET郑州工作组揭牌，国内首个国家级具身智能开源数据集社区落地郑州。与宇树科技共建「具身智能实训创新中心」和「具身智能产业学院」。中国具身智能市场规模2025年约9150亿元，预计2026年达10904亿元。郑州高新区已集聚400余家机器人产业链企业。",
                    "insight": "郑州借助龙头企业共建产业学院的模式值得借鉴。常州在新能源和智能制造领域可联合宁德时代、比亚迪、理想汽车等头部企业共建类似产业学院和实训中心，把「名企」资源在教育端和产业端双重落地。",
                    "source": "网易新闻"
                },
            ]
        },
        {
            "name": "改革举措",
            "items": [
                {
                    "title": "海南拟设1亿元「先投后股」资金池支持科技成果转化",
                    "date": "2026.7.3",
                    "summary": "海南发布2026年科技体制改革计划，全省「先投后股」资金池达1亿元，累计支持不少于20个项目。「先投后股」即财政资金先以补助形式投入支持研发，后期按约定转化为企业股权。同步推动省级财政科技专项体系优化改革，强化行业主管部门对科研项目的组织作用。",
                    "insight": "「先投后股」是破解成果转化「最初一公里」资金困境的有效模式。常州可在「三名联动」框架下试点——从产业基金中切出专项，以「先投后股」方式支持高校院所成果在常转化，实现财政资金循环使用和增值退出。",
                    "source": "海南中新网"
                },
                {
                    "title": "上海科学技术奖励大会：深化科技评价改革，推进「三评联动」",
                    "date": "2026.7.2",
                    "summary": "上海市委书记陈吉宁为褚君浩、陈赛娟颁发科技功臣奖。强调以国家战略需求为牵引深化科技体制改革，核心部署：加强科技计划全过程管理、深化项目经理团队制度建设、推进项目评审/人才评价/机构评估「三评联动」改革、为科研人员赋权增能减负松绑。",
                    "insight": "上海「三评联动」理念值得常州学习。常州可率先在市级科技计划中试行「分类评价」制度——基础研究看长周期、应用研究看转化效果、企业项目看市场效益，避免「一刀切」扼杀创新活力。",
                    "source": "上海市科委"
                },
                {
                    "title": "郑东新区构建硬科技项目全链条资金支持体系，破解成果转化「死亡谷」",
                    "date": "2026.7.2",
                    "summary": "郑东新区发布硬科技项目全链条资金支持：概念验证资金（1000万元/年）→产业研发联合基金（6000万元）→种子基金→天使投资→产业引导基金，形成从0到1再到N的全程资本护航。首批已拨付368万元支持6个硬科技项目，目标解决科研成果从实验室到产业化之间的融资断档。",
                    "insight": "常州的「三名联动」框架已具备「名园+名院+名企」三大要素，缺的正是贯穿全链条的资金机制。建议在中以常州创新园先行试点「概念验证+产业研发联合基金」模式，与常州产业投资基金联动，形成从验证到产业化的接力式资金供给。",
                    "source": "正观新闻"
                },
            ]
        },
    ]
    return generate_daily_from_data(sample_data)


def generate_weekly(output_path=None):
    """生成周报"""
    today = datetime.now()
    date_cn = today.strftime('%Y年%m月%d日')
    date_fn = today.strftime('%Y%m%d')
    issue, total = get_issue_numbers()

    doc = create_doc()
    add_header_block(doc, '创新常州·对标快讯（周报）', issue, total, date_cn)

    add_section_title(doc, '【本周综述】')
    p = doc.add_paragraph()
    run = p.add_run('本周核心主线为长三角国际科技创新中心建设加速推进，央地共建机制正式落地，沿沪宁产业创新带被明确为核心走廊，常州被列入高能级创新型城市建设对象。各地万亿城市纷纷加码科技创新投入，竞争态势加剧。')
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for section in ['各地科技委动态', '上海（长三角）国创中心资讯', '科创政策速览', '改革举措']:
        add_section_title(doc, f'【{section} · 本周要点】')
        p = doc.add_paragraph()
        run = p.add_run('（本周该板块主要条目汇总，由AI自动从日报聚合生成）')
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if output_path is None:
        output_path = PROJECT_DIR / 'weekly' / f'创新常州·对标快讯_周报_{date_fn}.docx'
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'[完成] 周报已生成: {output_path}')
    return output_path


def generate_monthly(output_path=None):
    """生成月报"""
    today = datetime.now()
    date_cn = today.strftime('%Y年%m月%d日')
    month_cn = today.strftime('%Y年%m月')
    issue, total = get_issue_numbers()

    doc = create_doc()
    add_header_block(doc, '创新常州·对标快讯（月报）', issue, total, date_cn)

    add_section_title(doc, '【月度综述】')
    p = doc.add_paragraph()
    run = p.add_run(f'{month_cn}，长三角国际科技创新中心建设持续推进，常州被列入高能级创新型城市建设对象。本月重点关注央地共建机制落地、沿沪宁产业创新带政策红利、各地万亿城市科技创新政策动态及科技体制改革进展。')
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for section in ['各地科技委动态', '上海（长三角）国创中心资讯', '科创政策速览', '改革举措']:
        add_section_title(doc, f'【{section} · 月度要点】')
        p = doc.add_paragraph()
        run = p.add_run('（本月该板块主要条目汇总，由AI自动从本月周报聚合生成）')
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    add_section_title(doc, '【月度战略建议】')
    suggestions = [
        '政策机遇窗口：常州被列入高能级创新型城市建设对象，应加快制定配套行动方案，争取省级资源倾斜。',
        '产业竞争态势：多个万亿城市加码同一赛道（新能源、具身智能、合成生物），常州需在技术深度和差异化定位上建立壁垒。',
        '改革红利：多地「先投后股」成果转化模式已成熟，常州可选取1-2个细分赛道先行先试，抢占改革试点先机。',
    ]
    for s in suggestions:
        p = doc.add_paragraph()
        run = p.add_run(f'• {s}')
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    if output_path is None:
        output_path = PROJECT_DIR / 'monthly' / f'创新常州·对标快讯_月报_{today.strftime("%Y%m")}.docx'
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'[完成] 月报已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='生成《创新常州·对标快讯》Word文档')
    parser.add_argument('--type', choices=['daily', 'weekly', 'monthly'], default='daily')
    parser.add_argument('--output', '-o', help='输出文件路径')
    parser.add_argument('--json', '-j', help='从JSON文件读取数据（用于数据驱动生成）')
    args = parser.parse_args()

    if args.json:
        with open(args.json, 'r', encoding='utf-8') as f:
            data = json.load(f)
        sections = data if isinstance(data, list) else data.get("sections", [])
        generate_daily_from_data(sections, args.output)
    elif args.type == 'daily':
        generate_daily(args.output)
    elif args.type == 'weekly':
        generate_weekly(args.output)
    elif args.type == 'monthly':
        generate_monthly(args.output)
